import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import datetime
import os
from PIL import Image
import PyPDF2
import logging
import time

# Configure logging
logging.basicConfig(level=logging.ERROR)
logger = logging.getLogger(__name__)

# --- CONSTANTS ---
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
API_TIMEOUT = 60  # seconds
MAX_RETRIES = 3

# --- 1. НАСТРОЙКА ИИ С ВАЛИДАЦИЕЙ ---
try:
    api_key = st.secrets.get("GOOGLE_API_KEY")
    if not api_key:
        st.error("❌ Google API Key не найден в secrets")
        st.stop()
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-flash-latest')
except Exception as e:
    st.error(f"❌ Ошибка инициализации API: {e}")
    logger.error(f"API initialization error: {e}")
    st.stop()

st.set_page_config(page_title="🐾Кошичкина Юстицыя🐾", page_icon="🐾", layout="wide")

# --- 2. SUPREME INTERFACE STYLING ---
st.markdown("""
    <style>
    .main { background-color: #0d1117; color: #c9d1d9; }
    h1, h2, h3 { color: #d4af37 !important; font-family: 'Georgia', serif; text-shadow: 2px 2px 4px rgba(0,0,0,0.5); }
    .stButton>button {
        background: linear-gradient(135deg, #d4af37 0%, #7a631d 100%);
        color: #fff !important; border-radius: 15px; font-weight: 900; height: 5em; width: 100%;
        border: 1px solid #d4af37; box-shadow: 0 4px 25px rgba(212, 175, 55, 0.3);
    }
    .stMetric { background-color: #1c2128; border: 1px solid #d4af37; padding: 20px; border-radius: 15px; }
    .stTextArea textarea { font-size: 16px !important; }
    </style>
    """, unsafe_allow_html=True)

# --- 3. CORE ANALYTICS ENGINES ---
def get_docx_text(file):
    """Extract text from DOCX file with error handling."""
    try:
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

def get_pdf_text(file):
    """Extract text from PDF file with error handling."""
    try:
        reader = PyPDF2.PdfReader(file)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        return ""

@st.cache_data(ttl=3600)
def load_all_standards():
    """Load all standard documents from standards folder with 1-hour cache."""
    folder = "standards"
    if not os.path.exists(folder):
        os.makedirs(folder)
    
    combined = ""
    try:
        for f in os.listdir(folder):
            if f.endswith('.docx'):
                try:
                    with open(os.path.join(folder, f), "rb") as file:
                        combined += f"\n[ЭТАЛОН: {f}]\n{get_docx_text(file)}\n"
                except Exception as e:
                    logger.warning(f"Ошибка загрузки стандарта {f}: {e}")
                    continue
    except Exception as e:
        logger.error(f"Error loading standards folder: {e}")
    
    return combined

# --- 4. SIDEBAR & TOOLS ---
with st.sidebar:
    st.title("🐾Кошичкина Юстицыя🐾")
    st.caption("Omniscient Final Edition | 2026")
    st.divider()
    
    all_standards = load_all_standards()
    st.success(f"📚 База эталонов: {len(os.listdir('standards')) if os.path.exists('standards') else 0}")
    
    files = st.file_uploader("📥 ВХОДЯЩИЕ (DOCX, PDF, ФОТО)", 
                            type=["docx", "pdf", "jpg", "jpeg", "png"], 
                            accept_multiple_files=True)
    
    # Validate file sizes
    if files:
        for f in files:
            if f.size > MAX_FILE_SIZE:
                st.error(f"❌ Файл '{f.name}' превышает лимит 50MB")
    
    st.metric("Базовая величина (2026)", "45,0 BYN")
    
    if st.button("♻️ Hard Reset System"):
        st.cache_data.clear()
        st.rerun()

# --- 5. WORKSPACE: ALL TASKS PRESERVED ---
col_l, col_r = st.columns([1, 1], gap="large")

with col_l:
    st.subheader("⚙️ Реестр задач (Без потерь)")
    task_type = st.selectbox("Выберите алгоритм:", [
        "🌟 OMNISCIENT: Максимальный контроль (КГС + Реклама + ГОСТ + НК)",
        "⚖️ СУДЕБНЫЙ ПРЕДИКТОР: Иск/Жалоба по КГС 2026 + Вероятность победы",
        "💰 КОНТРОЛЬ ЦЕН (713): Экспертиза надбавок и цен (МАРТ)",
        "🏦 БАНК-КОМПЛАЕНС: Проверка на 115-З (Блокировка счетов)",
        "📺 РЕКЛАМНЫЙ АУДИТ: Закон №225-З, Указ 95 и Реестр МАРТ",
        "🕵️ ДЕТЕКТОР МАНИПУЛЯЦИЙ: Поиск скрытых ловушек в чужом тексте",
        "🛡 АНТИ-ШТРАФ: Комплаенс ст. 33 НК (Дробление и выгода)",
        "📦 ПАКЕТНАЯ ГЕНЕРАЦИЯ: Договор + Протокол + Доверенность + Приказ",
        "🏛 СУДЕБНАЯ СТРАТЕГИЯ: Тезисы, контр-аргументы и КГС-тактика",
        "✉️ ГОС-ОТВЕТ: Ответы в ИМНС, ДФР, Госконтроль, Исполкомы",
        "🔒 НЕЙРО-СЕЙФ: Деперсонализация (Анонимизация образца)",
        "🔄 ВСТРЕЧНАЯ ПРОВЕРКА: Сверка допов с основным контрактом",
        "🔍 НЕЙРО-ВЫЧИТКА: Поиск конфликтов, битых ссылок и сленга",
        "⏰ ТАЙМ-МЕНЕДЖЕР: Процессуальные сроки КГС и давность",
        "📊 КАЛЬКУЛЯТОР 2.0: Госпошлина + ст. 366 ГК РБ",
        "🌐 ДВУЯЗЫЧНЫЙ КОНТРАКТ: Рус/Англ (Legal English)",
        "🚦 СВЕТОФОР РИСКОВ: Полный аудит (Red/Yellow/Green)",
        "🚢 ВЭД: Инструкция №37, Валютный контроль Нацбанка",
        "💡 IP-ПРАВО: Авторство, софт, исключительные права",
        "💼 ПЕРЕХОД ИП -> ООО: Полный пакет 'Бесшовно'",
        "📑 ИНСТРУКЦИЯ №4: Только проверка оформления по ГОСТу",
        "🛡 САНКЦИОННЫЙ ФИЛЬТР: Проверка сделки на ограничения",
        "📉 САММАРИ: Краткий отчет для Telegram",
        "СВОБОДНЫЙ ЮРИДИЧЕСКИЙ ЗАПРОС"
    ])
    user_input = st.text_area("Дополнительные условия (УНП, суммы, суть):", height=150)

with col_r:
    st.subheader("📋 Контекстные данные")
    images, text_ctx = [], ""
    if files:
        for f in files:
            # Skip files that exceed size limit
            if f.size > MAX_FILE_SIZE:
                continue
            
            ext = f.name.lower()
            try:
                if ext.endswith(('.jpg', '.jpeg', '.png')):
                    st.image(f, width=120)
                    images.append(Image.open(f))
                elif ext.endswith('.pdf'):
                    text_ctx += f"\n[PDF {f.name}]:\n{get_pdf_text(f)}\n"
                else:
                    text_ctx += f"\n[DOC {f.name}]:\n{get_docx_text(f)}\n"
            except Exception as e:
                logger.error(f"Error processing file {f.name}: {e}")
                st.warning(f"⚠️ Ошибка обработки файла {f.name}")
                continue
    else:
        st.info("Загрузите файлы сторон или шаблоны.")

# --- 6. FINAL EXECUTION ENGINE ---
if st.button("🚀 ЗАПУСТИТЬ ПРАВОВУЮ СИНГУЛЯРНОСТЬ"):
    if not files and not user_input:
        st.error("Ошибка: нет данных для анализа.")
    else:
        with st.spinner('Кошичка применяет все приоритеты и законы 2026...'):
            try:
                today_str = datetime.date.today().strftime("%d.%m.%Y")
                prompt = f"""
                ТЫ: 🐾 ULTIMATE PROMPT: KOSHICHKINA JUSTICE SUPREME 2026 (Full Version)
Роль:
Ты — «Кошичкина Юстицыя 25.0», специализированный ИИ-эксперт высшей категории по праву Республики Беларусь[...]\n                ЗАДАЧА: {task_type}.
                ДЕТАЛИ: {user_input}.
                ДАТА СЕГОДНЯ: {today_str}.
                ...
                БАЗА ЭТАЛОНОВ: {all_standards}
                КОНТЕКСТ ФАЙЛОВ: {text_ctx}
                """
                
                # Retry logic with exponential backoff
                res = None
                for attempt in range(MAX_RETRIES):
                    try:
                        res = model.generate_content([prompt] + images, timeout=API_TIMEOUT)
                        break
                    except Exception as e:
                        if attempt < MAX_RETRIES - 1:
                            wait_time = 2 ** attempt
                            st.warning(f"⚠️ Попытка {attempt + 1} не удалась, повтор через {wait_time}с...")
                            time.sleep(wait_time)
                            logger.warning(f"API attempt {attempt + 1} failed: {e}")
                        else:
                            raise
                
                if not res:
                    raise Exception("API did not return a response")
                
                st.divider()
                st.subheader("🏁 Результат:")
                st.markdown(res.text)

                # --- ADVANCED WORD EXPORT (TIMES 14 + GOST) ---
                doc_out = Document()
                
                # Поля по ГОСТ: Левое 30мм, Правое 10мм, Верх/Низ 20мм
                sec = doc_out.sections[0]
                sec.left_margin = Cm(3.0)
                sec.right_margin = Cm(1.0)
                sec.top_margin = Cm(2.0)
                sec.bottom_margin = Cm(2.0)

                style = doc_out.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                
                # Настройка кириллицы (FIXED: moved outside loop)
                rfonts = style.element.rPr.rFonts
                for t in ['ascii', 'hAnsi', 'eastAsia']:
                    rfonts.set(qn(f'w:{t}'), 'Times New Roman')
                
                # Add heading (FIXED: now outside the font configuration loop)
                doc_out.add_heading(f'{task_type}', 0)
                
                for line in res.text.split('\n'):
                    p = doc_out.add_paragraph(line)
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Cm(1.25)
                    p.paragraph_format.line_spacing = 1.0

                b = BytesIO()
                doc_out.save(b)
                b.seek(0)
                st.download_button(
                    "📥 СКАЧАТЬ DOCX (ГОСТ РБ: Times 14)",
                    b,
                    f"Koshichkina_Justice_Final_{datetime.date.today()}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"❌ Критический сбой: {e}")
                logger.error(f"Critical error: {e}", exc_info=True)

st.divider()
st.markdown("<center style='opacity:0.3;'>Кошичкина Юстицыя 25.0 Omniscient Supreme • All Priorities Integrated • 2026</center>", unsafe_allow_html=True)